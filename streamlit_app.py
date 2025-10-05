import os
import datetime
import streamlit as st
from openai import OpenAI, APIError
from docx import Document
import tempfile
import json
import pandas as pd
import re
from typing import Dict, Any, Union

# 루브릭 기준 정의
RUBRIC_CRITERIA = {
    "내용논리성": {
        "4점 (우수)": "주장이 명확하고 충분한 근거가 체계적으로 제시됨. 논리적 연결이 자연스럽고 설득력이 있음",
        "3점 (보통)": "주장은 명확하나 근거가 부분적으로 부족하거나 논리적 연결이 일부 어색함",
        "2점 (미흡)": "주장이 다소 모호하고 근거가 약함. 논리적 연결에 문제가 있음",
        "1점 (부족)": "주장과 근거가 불분명하고 논리적 흐름이 부자연스러움"
    },
    "구성체계성": {
        "4점 (우수)": "서론-본론-결론 구조가 명확하고 문단 간 연결이 자연스러움. 응집성과 일관성이 뛰어남",
        "3점 (보통)": "전체 구조는 갖추었으나 문단 간 연결이 부분적으로 어색함",
        "2점 (미흡)": "구조가 불분명하거나 문단 간 흐름이 부자연스러움", 
        "1점 (부족)": "전체 구성이 체계적이지 않고 일관성이 부족함"
    },
    "문법어휘정확성": {
        "4점 (우수)": "문법적 오류가 거의 없고 어휘 사용이 적절하며 문장 구조가 다양함",
        "3점 (보통)": "문법적 오류가 약간 있으나 의미 전달에 큰 문제없음. 어휘 사용이 대체로 적절함",
        "2점 (미흡)": "문법적 오류가 자주 발생하고 어휘 선택이 부적절한 경우가 있음",
        "1점 (부족)": "문법적 오류가 많고 어휘 사용이 부정확하여 의미 전달에 어려움이 있음"
    }
}

# 환경 설정
try:
    OPENAI_KEY = st.secrets["openai"]["api_key"]
except KeyError:
    OPENAI_KEY = ""

OPENAI_OK = bool(OPENAI_KEY)

client = None
if OPENAI_OK:
    try:
        client = OpenAI(api_key=OPENAI_KEY)
    except Exception as e:
        st.error(f"OpenAI 초기화 실패: {e}")
        OPENAI_OK = False

# 누락된 함수 추가 - 문제 1 해결
def display_rubric():
    """루브릭 기준 표시"""
    st.markdown("### 평가 기준 (루브릭)")
    st.markdown("글을 쓰기 전에 평가 기준을 확인해보세요!")
    
    for category, criteria in RUBRIC_CRITERIA.items():
        with st.expander(f"{category} 평가 기준", expanded=False):
            for score, description in criteria.items():
                st.markdown(f"**{score}**: {description}")

# 새로운 유틸리티 함수 추가
def parse_gpt_json_response(response_text: str) -> dict:
    """GPT 응답에서 JSON 블록을 추출하고 파싱"""
    try:
        # ```json 블록에서 JSON 추출
        if "```json" in response_text:
            json_match = re.search(r'```json\s*\n(.*?)\n```', response_text, re.DOTALL)
            if json_match:
                json_str = json_match.group(1).strip()
            else:
                json_str = response_text.replace('```json\n', '').replace('\n```', '').strip()
        else:
            json_str = response_text.strip()
        
        # JSON 파싱 시도 (더 유연하게)
        return json.loads(json_str)
    except json.JSONDecodeError as e:
        # JSON 파싱 실패 시 원본 텍스트를 포함한 오류 정보 반환
        return {
            "error": f"JSON 파싱 실패: {e}",
            "raw_response": response_text
        }
    except Exception as e:
        return {
            "error": f"응답 처리 실패: {e}",
            "raw_response": response_text
        }

def format_analysis_for_display(analysis_data: Union[dict, str], analysis_type: str = "analysis") -> dict:
    """분석 데이터를 표시용으로 포맷팅"""
    if not analysis_data:
        return {"error": "데이터 없음"}
    
    if isinstance(analysis_data, dict) and "error" not in analysis_data:
        return analysis_data
    
    if isinstance(analysis_data, dict) and analysis_type in analysis_data:
        return parse_gpt_json_response(analysis_data[analysis_type])
    
    return analysis_data

def format_docx_section(doc, title: str, data: dict):
    """DOCX 문서에 특정 분석 섹션을 포맷팅하여 추가"""
    doc.add_heading(title, level=2)
    
    if "error" in data:
        doc.add_paragraph(f"오류: {data['error']}")
        return
        
    for key, value in data.items():
        if isinstance(value, dict):
            doc.add_paragraph(f"• {key}: {value.get('점수', 'N/A')}/4점")
            if '근거' in value:
                doc.add_paragraph(f"    - 근거: {value['근거']}")
            if '개선제안' in value:
                doc.add_paragraph(f"    - 개선 제안: {value['개선제안']}")
        elif isinstance(value, list):
            doc.add_paragraph(f"• {key}: {', '.join(value)}")
        else:
            doc.add_paragraph(f"• {key}: {value}")
    doc.add_paragraph("")

def get_paragraph_feedback(text: str, paragraph_type: str, context: dict = None) -> dict:
    """문단별 즉시 피드백 제공"""
    if not OPENAI_OK or client is None:
        return {"error": "API 오류"}
    
    if not text.strip():
        return {"error": "입력된 텍스트가 없습니다"}
    
    context_info = ""
    if context:
        if context.get("summary1") and context.get("summary2"):
            context_info = f"""
            참고용 기사 요약:
            기사1: {context['summary1'][:200]}...
            기사2: {context['summary2'][:200]}...
            """
    
    prompt = f"""
    다음 {paragraph_type} 문단을 분석하고 즉시 개선할 수 있는 구체적인 피드백을 제공하세요.

    {context_info}

    문단 내용: {text}

    응답은 반드시 다음 JSON 형식으로만 제공하세요:
    {{
        "강점": ["구체적 강점1", "구체적 강점2"],
        "개선점": ["개선사항1", "개선사항2"],
        "구체적제안": "즉시 적용 가능한 수정 제안",
        "추천점수": 1~4 사이의 정수값
    }}
    """
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=500
        )
        
        result = response.choices[0].message.content.strip()
        return parse_gpt_json_response(result)
    except APIError as e:
        return {"error": f"피드백 생성 실패: OpenAI API 오류 - {e}"}
    except Exception as e:
        return {"error": f"피드백 생성 실패: {e}"}

def summarize_paragraph_feedback(paragraph_feedback: dict) -> str:
    """문단별 피드백을 요약"""
    if not paragraph_feedback:
        return "문단별 피드백 없음"
    
    summary = "문단별 학습 과정 요약:\n"
    for section, feedback in paragraph_feedback.items():
        if isinstance(feedback, dict):
            if "error" not in feedback:
                summary += f"- {section}: 추천점수 {feedback.get('추천점수', 'N/A')}점\n"
                summary += f"  주요 개선제안: {feedback.get('구체적제안', 'N/A')}\n"
            else:
                summary += f"- {section}: 피드백 없음\n"
    return summary

def summarize_text(text: str) -> str:
    """기사 내용을 영어로 5-10문장으로 요약"""
    if not OPENAI_OK or client is None:
        return "요약 불가: API 오류"
    if not text.strip():
        return "요약 불가: 입력된 텍스트가 없습니다."
    
    prompt = f"다음 기사 내용을 영어로 다섯 문장 이상, 열문장 이하로 요약해줘:\n\n{text}"
    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=800
        )
        return response.choices[0].message.content.strip()
    except APIError as e:
        return f"요약 실패: OpenAI API 오류 - {e}"
    except Exception as e:
        return f"요약 실패: {e}"

def analyze_tone_and_stance(text: str) -> dict:
    """논조 및 입장 분석 - 점수화된 논조 포함"""
    if not OPENAI_OK or client is None:
        return {"error": "API 오류"}
    
    prompt = f"""
    다음 기사의 논조와 입장을 분석해주세요.
    
    응답은 반드시 다음 JSON 형식으로만 제공하고, 다른 텍스트는 포함하지 마세요:
    
    {{
        "논조분류": "positive/neutral/negative",
        "논조점수": -3~3 사이의 정수값,
        "주요논점": ["논점1", "논점2", "논점3"],
        "감정적언어": ["예시1", "예시2", "예시3"],
        "신뢰도점수": 1~10 사이의 정수값,
        "객관성점수": 1~10 사이의 정수값
    }}
    
    기사: {text}
    """
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=600
        )
        
        result = response.choices[0].message.content.strip()
        return parse_gpt_json_response(result)
    except APIError as e:
        return {"error": f"분석 실패: OpenAI API 오류 - {e}"}
    except Exception as e:
        return {"error": f"분석 실패: {e}"}

def evaluate_writing_rubric(text: str) -> dict:
    """영어 표현 능력 루브릭 평가"""
    if not OPENAI_OK or client is None:
        return {"error": "API 오류"}
    
    prompt = f"""
    다음 영어 텍스트를 구체적인 루브릭 기준으로 평가해주세요.

    **평가 영역 및 기준:**

    **1. 내용 논리성 (Content Logic) - 1~4점**
    **2. 구성 체계성 (Organization) - 1~4점**
    **3. 문법·어휘 정확성 (Language Accuracy) - 1~4점**

    응답은 반드시 다음 JSON 형식으로만 제공하고, 다른 텍스트는 포함하지 마세요:

    {{
        "내용논리성": {{
            "점수": 1~4 사이의 정수값,
            "근거": "구체적 평가 근거"
        }},
        "구성체계성": {{
            "점수": 1~4 사이의 정수값, 
            "근거": "구체적 평가 근거"
        }},
        "문법어휘정확성": {{
            "점수": 1~4 사이의 정수값,
            "근거": "구체적 평가 근거"
        }},
        "총점": "12점 만점 중 X점",
        "종합평가": "전체적인 평가 및 개선 제안"
    }}

    평가 대상 텍스트: {text}
    """
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
            max_tokens=800
        )
        
        result = response.choices[0].message.content.strip()
        return parse_gpt_json_response(result)
    except APIError as e:
        return {"error": f"평가 실패: OpenAI API 오류 - {e}"}
    except Exception as e:
        return {"error": f"평가 실패: {e}"}

def assess_problem_solving(reflection_text: str) -> dict:
    """문제해결 역량 평가"""
    if not OPENAI_OK or client is None:
        return {"error": "API 오류"}
    
    if not reflection_text or len(reflection_text.strip()) < 10:
        return {
            "assessment": "성찰 내용이 너무 짧아 평가하기에 정보가 부족합니다. 문제 해결 과정에 대한 구체적인 설명을 포함해 주세요."
        }
    
    prompt = f"""
    다음 학습자의 성찰 내용을 바탕으로 문제해결 역량을 평가해주세요:
    
    응답은 반드시 다음 JSON 형식으로만 제공하고, 다른 텍스트는 포함하지 마세요:
    
    {{
        "문제이해": {{
            "점수": 1~5 사이의 정수값,
            "개선제안": "구체적인 개선 제안"
        }},
        "분석적사고": {{
            "점수": 1~5 사이의 정수값,
            "개선제안": "구체적인 개선 제안"
        }},
        "대안발견및기획": {{
            "점수": 1~5 사이의 정수값,
            "개선제안": "구체적인 개선 제안"
        }},
        "의사소통": {{
            "점수": 1~5 사이의 정수값,
            "개선제안": "구체적인 개선 제안"
        }},
        "총점": "20점 만점 중 X점",
        "종합평가": "전체적인 평가 요약"
    }}
    
    성찰 내용: {reflection_text}
    """
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
            max_tokens=800
        )
        
        result = response.choices[0].message.content.strip()
        return parse_gpt_json_response(result)
    except APIError as e:
        return {"error": f"평가 실패: OpenAI API 오류 - {e}"}
    except Exception as e:
        return {"error": f"평가 실패: {e}"}

def display_emotional_words(analysis1: dict, analysis2: dict) -> None:
    """감정적 언어 시각화"""
    st.markdown("#### 감정적 표현 비교")
    
    analysis1 = format_analysis_for_display(analysis1, "analysis")
    analysis2 = format_analysis_for_display(analysis2, "analysis")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("**기사 1**")
        if "감정적언어" in analysis1 and analysis1["감정적언어"]:
            words = analysis1["감정적언어"]
            word_html = ""
            for i, word in enumerate(words):
                size = 20 - i*2
                color = ["#ff6b6b", "#4ecdc4", "#45b7d1", "#96ceb4", "#feca57"][i % 5]
                word_html += f'<span style="font-size:{size}px; color:{color}; margin:5px; font-weight:bold;">{word}</span> '
            
            st.markdown(f'<div style="line-height:2;">{word_html}</div>', unsafe_allow_html=True)
        else:
            st.info("감정적 표현이 감지되지 않았습니다.")
    
    with col2:
        st.markdown("**기사 2**")
        if "감정적언어" in analysis2 and analysis2["감정적언어"]:
            words = analysis2["감정적언어"]
            word_html = ""
            for i, word in enumerate(words):
                size = 20 - i*2
                color = ["#ff6b6b", "#4ecdc4", "#45b7d1", "#96ceb4", "#feca57"][i % 5]
                word_html += f'<span style="font-size:{size}px; color:{color}; margin:5px; font-weight:bold;">{word}</span> '
            
            st.markdown(f'<div style="line-height:2;">{word_html}</div>', unsafe_allow_html=True)
        else:
            st.info("감정적 표현이 감지되지 않았습니다.")

def create_simple_gauge(value: int, title: str) -> None:
    """간단한 게이지 시각화"""
    percentage = ((value + 3) / 6) * 100
    
    if value <= -2:
        color = "#ff4757"
        emoji = "😞"
    elif value <= -1:
        color = "#ffa502"
        emoji = "😕"
    elif value == 0:
        color = "#747d8c"
        emoji = "😐"
    elif value <= 1:
        color = "#2ed573"
        emoji = "🙂"
    else:
        color = "#5352ed"
        emoji = "😊"
    
    gauge_html = f"""
    <div style="text-align: center; margin: 20px; padding: 20px; border-radius: 10px; background-color: #f8f9fa;">
        <h4 style="margin-bottom: 15px;">{title}</h4>
        <div style="font-size: 30px; margin-bottom: 10px;">{emoji}</div>
        <div style="width: 200px; height: 20px; background-color: #e1e8ed; border-radius: 10px; margin: 0 auto; position: relative;">
            <div style="width: {percentage}%; height: 100%; background-color: {color}; border-radius: 10px;"></div>
        </div>
        <p style="margin-top: 10px; font-weight: bold; color: {color}; font-size: 18px;">{value}점</p>
    </div>
    """
    st.markdown(gauge_html, unsafe_allow_html=True)

def create_enhanced_comparison_chart(analysis1: dict, analysis2: dict) -> None:
    """개선된 논조 비교 차트"""
    analysis1 = format_analysis_for_display(analysis1, "analysis")
    analysis2 = format_analysis_for_display(analysis2, "analysis")
    
    if "error" in analysis1 or "error" in analysis2:
        st.error("논조 분석 데이터가 부족하여 차트를 생성할 수 없습니다.")
        return
    
    st.markdown("#### 논조 점수 비교")
    
    col1, col2 = st.columns(2)
    
    with col1:
        create_simple_gauge(analysis1.get('논조점수', 0), "기사 1 논조")
    with col2:
        create_simple_gauge(analysis2.get('논조점수', 0), "기사 2 논조")
    
    st.markdown("#### 신뢰도 및 객관성")
    
    trust1 = analysis1.get('신뢰도점수', 5)
    trust2 = analysis2.get('신뢰도점수', 5)
    obj1 = analysis1.get('객관성점수', 5)
    obj2 = analysis2.get('객관성점수', 5)
    
    metrics_df = pd.DataFrame({
        '지표': ['신뢰도', '객관성'],
        '기사1': [trust1, obj1],
        '기사2': [trust2, obj2]
    })
    st.bar_chart(metrics_df.set_index('지표'))

def gpt_feedback(korean_text: str, reflection_text: str = "") -> str:
    """한국어 작문에 대한 한국어 피드백 제공 (기존 함수 유지)"""
    if not OPENAI_OK or client is None:
        return "GPT 사용을 위한 OpenAI API 키가 설정되지 않았거나 문제가 있습니다."
    if not korean_text.strip():
        return "피드백할 텍스트가 없습니다."
        
    reflection_info = ""
    if reflection_text:
        reflection_info = f"""
        학습자가 남긴 성찰 내용:
        {reflection_text}
        
        **이 성찰 내용을 참고하여, 학습자가 어려움을 느꼈거나 개선하고 싶다고 언급한 부분을 중심으로 피드백을 강화해주세요.**
        """

    prompt = f"""
    당신은 한국인 학습자를 위한 글쓰기 지도교사입니다. 다음 비교 설명문을 평가하고 건설적인 피드백을 한국어로 제공하세요.

    다음 루브릭 기준으로 평가해주세요:

    **1. 내용 논리성 (Content Logic)** - 주장의 명확성, 근거 제시 충분성, 논리적 연결, 문제 상황 분석 깊이

    **2. 구성 체계성 (Organization)**
    - 서론-본론-결론 구조, 문단 간 연결과 흐름, 응집성과 일관성

    **3. 문법·어휘 정확성 (Language Accuracy)**
    - 문법적 정확성, 문장 구조의 다양성, 어휘 선택의 적절성

    각 영역별로 구체적인 피드백과 3-5개의 개선 제안을 제공해주세요.

    {reflection_info}

    평가 대상 글:
    {korean_text}
    """

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "당신은 글쓰기 지도교사입니다."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.2,
            max_tokens=1200
        )
        return response.choices[0].message.content.strip()
    except APIError as e:
        return f"OpenAI API 오류: {e}"
    except Exception as e:
        return f"GPT 호출 오류: {e}"

def translate_to_korean(text: str) -> str:
    """영문 텍스트를 한국어로 번역"""
    if not OPENAI_OK or client is None:
        return "번역 불가: API 오류"
    if "요약 실패" in text or "요약 불가" in text:
        return "원본 요약이 없어 번역할 수 없습니다."
    if not text.strip():
        return "번역 불가: 입력된 텍스트가 없습니다."

    prompt = f"다음 영문 텍스트를 자연스러운 한국어로 번역해줘:\n\n{text}"
    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=1000
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"번역 실패: {e}"

def translate_to_english(text: str) -> str:
    """한국어 텍스트를 영어로 번역"""
    if not OPENAI_OK or client is None:
        return "번역 불가: API 오류"
    if not text.strip():
        return "번역 불가: 입력된 텍스트가 없습니다."

    prompt = f"다음 한국어 텍스트를 자연스러운 영어로 번역해줘:\n\n{text}"
    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=1200
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"번역 실패: {e}"

def create_docx_content(text: str, analysis_data: Dict[str, Any]) -> bytes:
    """텍스트와 분석 데이터를 DOCX 파일로 변환하여 바이트 데이터 반환"""
    doc = Document()
    doc.add_heading('News Comparison Analysis', 0)
    doc.add_paragraph(f"작성일: {datetime.datetime.now().strftime('%Y년 %m월 %d일 %H:%M')}")
    doc.add_paragraph("")
    
    if analysis_data:
        doc.add_heading('분석 요약', level=1)
        
        for key, value in analysis_data.items():
            if isinstance(value, dict) and "error" not in value:
                format_docx_section(doc, key, value)
            elif isinstance(value, str):
                doc.add_heading(key, level=2)
                doc.add_paragraph(value)
                doc.add_paragraph("")
            else:
                doc.add_heading(key, level=2)
                doc.add_paragraph("데이터 없음 또는 오류.")
                doc.add_paragraph("")

    doc.add_heading('작성된 설명문', level=1)
    for line in text.splitlines():
        if line.strip():
            doc.add_paragraph(line)
    
    with tempfile.NamedTemporaryFile() as tmp:
        doc.save(tmp.name)
        tmp.seek(0)
        return tmp.read()

def read_uploaded_file(uploaded_file) -> str:
    """업로드된 파일을 읽어서 텍스트 반환"""
    try:
        try:
            content = uploaded_file.read().decode('utf-8')
        except UnicodeDecodeError:
            uploaded_file.seek(0)
            content = uploaded_file.read().decode('euc-kr')
        return content
    except Exception as e:
        return f"파일 읽기 오류: {e}"

# Streamlit 앱 설정
st.set_page_config(
    page_title="News Comparison Assistant", 
    layout="wide",
    initial_sidebar_state="collapsed"
)

# 문제 2 해결: 누락된 세션 상태 키들 추가
if "stage" not in st.session_state:
    st.session_state.update({
        "stage": "input",
        "article1": "", "article2": "",
        "uploaded_file1_content": "", "uploaded_file2_content": "",
        "summary1": "", "summary2": "",
        "summary1_kr": "", "summary2_kr": "",
        "tone_analysis1": {}, "tone_analysis2": {},
        "draft": "", "feedback": "",
        "writing_evaluation": {},
        "problem_solving_score": {},
        "reflection_log": [],
        "final_text": "",
        "paragraph_feedback": {},
        # 누락된 키들 추가
        "intro_input": "",
        "body1_input": "",
        "body2_input": "",
        "compare_input": "",
        "conclusion_input": ""
    })

st.title("News Comparison and Writing Assistant")

if not OPENAI_OK:
    st.warning("OpenAI API 키가 설정되지 않았거나 문제가 있습니다. 요약 및 피드백 기능이 비활성화됩니다.")

progress_stages = ["input", "analysis", "draft", "feedback", "final"]
current_stage_idx = progress_stages.index(st.session_state.stage)
progress = (current_stage_idx + 1) / len(progress_stages)

st.progress(progress)
stage_names = ["기사 입력", "논조 분석", "초안 작성", "AI 피드백", "최종 완성"]
st.caption(f"현재 단계: {stage_names[current_stage_idx]} ({current_stage_idx + 1}/{len(progress_stages)})")

if st.session_state.stage in ["draft", "feedback"]:
    display_rubric()
    st.markdown("---")

if st.session_state.stage == "input":
    st.subheader("1단계. 기사 본문 입력 (파일 업로드)")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("**기사 1**")
        uploaded_file1 = st.file_uploader(
            "첫 번째 기사 파일 업로드 (.txt)", 
            type=['txt'], 
            key="file1",
            help="UTF-8 또는 EUC-KR 인코딩된 텍스트 파일을 드래그하여 업로드하세요 (최대 10MB)"
        )
        
        if uploaded_file1:
            st.session_state.uploaded_file1_content = read_uploaded_file(uploaded_file1)
            st.session_state.article1 = st.session_state.uploaded_file1_content
            if st.session_state.uploaded_file1_content.startswith("파일 읽기 오류"):
                st.error(st.session_state.uploaded_file1_content)
    
    with col2:
        st.markdown("**기사 2**")
        uploaded_file2 = st.file_uploader(
            "두 번째 기사 파일 업로드 (.txt)", 
            type=['txt'], 
            key="file2",
            help="UTF-8 또는 EUC-KR 인코딩된 텍스트 파일을 드래그하여 업로드하세요 (최대 10MB)"
        )
        
        if uploaded_file2:
            st.session_state.uploaded_file2_content = read_uploaded_file(uploaded_file2)
            st.session_state.article2 = st.session_state.uploaded_file2_content
            if st.session_state.uploaded_file2_content.startswith("파일 읽기 오류"):
                st.error(st.session_state.uploaded_file2_content)
    
    st.markdown("---")
    
    col_btn1, col_btn2 = st.columns([1, 1])
    with col_btn2:
        overall_error_placeholder = st.empty()
        if st.button("다음 단계 → (분석 시작)", type="primary", use_container_width=True):
            is_valid = True
            if not st.session_state.article1.strip():
                st.error("기사 1 파일을 업로드해주세요.")
                is_valid = False

            if not st.session_state.article2.strip():
                st.error("기사 2 파일을 업로드해주세요.")
                is_valid = False

            if is_valid:
                st.session_state.stage = "analysis"
                st.rerun()
            else:
                st.warning("모든 필수 입력 필드를 채워주세요.")

elif st.session_state.stage == "analysis":
    st.subheader("2단계. 논조 분석 및 요약")
    
    if not st.session_state.get("summary1") or not st.session_state.get("tone_analysis1"):
        with st.spinner("기사 분석 및 요약을 생성하고 있습니다..."):
            st.info("기사 1 요약 중...", icon="📝")
            st.session_state.summary1 = summarize_text(st.session_state.article1)
            st.info("기사 2 요약 중...", icon="📝")
            st.session_state.summary2 = summarize_text(st.session_state.article2)
            
            st.info("기사 1 논조 분석 중...", icon="🔎")
            st.session_state.tone_analysis1 = analyze_tone_and_stance(st.session_state.article1)
            st.info("기사 2 논조 분석 중...", icon="🔎")
            st.session_state.tone_analysis2 = analyze_tone_and_stance(st.session_state.article2)
            
            st.info("요약문 번역 중...", icon="🌐")
            st.session_state.summary1_kr = translate_to_korean(st.session_state.summary1)
            st.session_state.summary2_kr = translate_to_korean(st.session_state.summary2)
        st.success("모든 분석이 완료되었습니다!")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("#### 기사 1 분석")
        with st.expander("요약 (영어/한국어)", expanded=True):
            st.info(f"**[English]**\n{st.session_state.summary1}")
            st.success(f"**[한국어]**\n{st.session_state.summary1_kr}")
        
        with st.expander("논조 분석", expanded=True):
            analysis1 = format_analysis_for_display(st.session_state.tone_analysis1, "analysis")
            if "error" not in analysis1:
                st.markdown(f"**논조**: {analysis1.get('논조분류', 'N/A')} ({analysis1.get('논조점수', 0)}점)")
                st.markdown(f"**신뢰도**: {analysis1.get('신뢰도점수', 0)}/10점")
                st.markdown(f"**객관성**: {analysis1.get('객관성점수', 0)}/10점")
                
                if analysis1.get('주요논점'):
                    st.markdown("**주요 논점**:")
                    for i, point in enumerate(analysis1.get('주요논점', []), 1):
                        st.markdown(f"  {i}. {point}")
            else:
                st.error(analysis1.get("error", "분석 오류"))
    
    with col2:
        st.markdown("#### 기사 2 분석")
        with st.expander("요약 (영어/한국어)", expanded=True):
            st.info(f"**[English]**\n{st.session_state.summary2}")
            st.success(f"**[한국어]**\n{st.session_state.summary2_kr}")
        
        with st.expander("논조 분석", expanded=True):
            analysis2 = format_analysis_for_display(st.session_state.tone_analysis2, "analysis")
            if "error" not in analysis2:
                st.markdown(f"**논조**: {analysis2.get('논조분류', 'N/A')} ({analysis2.get('논조점수', 0)}점)")
                st.markdown(f"**신뢰도**: {analysis2.get('신뢰도점수', 0)}/10점")
                st.markdown(f"**객관성**: {analysis2.get('객관성점수', 0)}/10점")
                
                if analysis2.get('주요논점'):
                    st.markdown("**주요 논점**:")
                    for i, point in enumerate(analysis2.get('주요논점', []), 1):
                        st.markdown(f"  {i}. {point}")
            else:
                st.error(analysis2.get("error", "분석 오류"))
    
    st.markdown("---")
    create_enhanced_comparison_chart(st.session_state.tone_analysis1, st.session_state.tone_analysis2)
    
    st.markdown("---")
    display_emotional_words(st.session_state.tone_analysis1, st.session_state.tone_analysis2)
    
    st.markdown("---")
    st.markdown("#### 분석 성찰")
    reflection_error_placeholder = st.empty()
    reflection = st.text_area(
        "두 기사의 차이점과 공통점, 그리고 각각의 논조에 대한 당신의 생각을 적어보세요:",
        height=100,
        key="analysis_reflection"
    )
    
    col_btn1, col_btn2 = st.columns([1, 1])
    with col_btn1:
        if st.button("← 이전 단계", use_container_width=True):
            st.session_state.stage = "input"
            st.rerun()
    
    with col_btn2:
        if st.button("초안 작성 단계로 →", type="primary", use_container_width=True):
            if not reflection.strip():
                reflection_error_placeholder.error("분석 성찰 내용을 입력해주세요.")
            else:
                reflection_error_placeholder.empty()
                st.session_state.reflection_log.append({
                    "stage": "analysis",
                    "content": reflection,
                    "timestamp": datetime.datetime.now()
                })
                st.session_state.stage = "draft"
                st.rerun()

elif st.session_state.stage == "draft":
    st.subheader("3단계. 비교 설명문 초안 작성")

    # 문제 4 해결: 개선된 문단 입력 함수
    def paragraph_input_with_guide(title, key, guide_title, guide_lines, summary_text=None, hint_key=None, hint_prompt=None):
        col1, col2 = st.columns([2, 1])
        with col1:
            st.subheader(title)
            
            # 현재 값 가져오기
            current_value = st.session_state.get(key, "")
            
            # 고유한 키로 text_area 생성
            user_input = st.text_area(
                "", 
                value=current_value, 
                key=f"input_{key}_{st.session_state.stage}", 
                height=160
            )
            
            # 세션 상태 업데이트
            st.session_state[key] = user_input
            
            col1_1, col1_2 = st.columns([1, 1])
            with col1_1:
                if st.button(f"{title.split(' ')[0]} 완료", key=f"{key}_complete"):
                    if user_input.strip():
                        context = {
                            "summary1": st.session_state.get("summary1"),
                            "summary2": st.session_state.get("summary2")
                        }
                        feedback = get_paragraph_feedback(user_input, title, context)
                        st.session_state.paragraph_feedback[key] = feedback
                        st.success("피드백을 확인하세요!")
                    else:
                        st.error("문단을 먼저 작성해주세요.")
            
            with col1_2:
                if key in st.session_state.paragraph_feedback:
                    feedback = st.session_state.paragraph_feedback[key]
                    if "error" not in feedback:
                        score = feedback.get('추천점수', 'N/A')
                        st.metric("추천점수", f"{score}/4점")
            
            if key in st.session_state.paragraph_feedback:
                feedback = st.session_state.paragraph_feedback[key]
                if "error" not in feedback:
                    with st.expander("문단별 피드백", expanded=True):
                        if feedback.get('강점'):
                            st.markdown("**강점:**")
                            for strength in feedback['강점']:
                                st.markdown(f"- {strength}")
                        
                        if feedback.get('개선점'):
                            st.markdown("**개선점:**")
                            for improvement in feedback['개선점']:
                                st.markdown(f"- {improvement}")
                        
                        if feedback.get('구체적제안'):
                            st.markdown("**구체적 제안:**")
                            st.info(feedback['구체적제안'])
                else:
                    st.error(feedback.get("error", "피드백 오류"))
        
        with col2:
            st.markdown(f"#### {guide_title}")
            for line in guide_lines:
                st.markdown(f"- {line}")
            
            if summary_text:
                st.markdown("#### 관련 기사 요약")
                summary_en = summary_text
                summary_kr_key = "summary1_kr" if summary_en == st.session_state.get("summary1") else "summary2_kr"
                summary_kr = st.session_state.get(summary_kr_key, "번역 없음")
                
                with st.expander("요약문 보기", expanded=True):
                    st.info(f"**[English]**\n{summary_en}")
                    st.success(f"**[한국어]**\n{summary_kr}")

            if hint_key and hint_prompt and OPENAI_OK:
                if f"{hint_key}_hint" not in st.session_state:
                    st.session_state[f"{hint_key}_hint"] = ""
                if st.button(f"AI 힌트 받기", key=f"{hint_key}_btn"):
                    with st.spinner("AI 힌트를 생성하는 중입니다..."):
                        try:
                            hint_response = client.chat.completions.create(
                                model="gpt-4o",
                                messages=[{"role": "user", "content": hint_prompt}],
                                temperature=0.5,
                                max_tokens=300
                            )
                            st.session_state[f"{hint_key}_hint"] = hint_response.choices[0].message.content.strip()
                        except Exception as e:
                            st.session_state[f"{hint_key}_hint"] = f"힌트 생성 실패: {e}"
                if st.session_state[f"{hint_key}_hint"]:
                    st.markdown("#### AI 힌트")
                    st.success(st.session_state[f"{hint_key}_hint"])
        
        return user_input

    intro = paragraph_input_with_guide(
        "서론", "intro_input", "비교 주제 소개", [
            "비교할 두 기사 간단히 소개",
            "글의 목적, 문제 제기",
            "두 관점 간 차이에 대한 암시"
        ],
        hint_key="intro", hint_prompt="비교 설명문의 서론을 쓰기 위한 문장 구성 힌트를 3개 제시해줘. (한국어)"
    )

    body1 = paragraph_input_with_guide(
        "본론 - 기사 1 설명", "body1_input", "기사 1 요약", [
            "기사 1의 주장과 근거 요약",
            "자료, 사례, 강조점 기술"
        ],
        summary_text=st.session_state.get("summary1"),
        hint_key="body1", hint_prompt="첫 번째 기사 내용을 요약하는 문단 작성에 쓸 수 있는 문장 예시 3개를 제시해줘. (한국어)"
    )

    body2 = paragraph_input_with_guide(
        "본론 - 기사 2 설명", "body2_input", "기사 2 요약", [
            "기사 2의 주요 내용 요약",
            "기사 1과 비교했을 때의 특징 언급"
        ],
        summary_text=st.session_state.get("summary2"),
        hint_key="body2", hint_prompt="두 번째 기사 내용을 요약하며 비교하는 문단을 쓰기 위한 문장 예시 3개를 제시해줘. (한국어)"
    )

    compare = paragraph_input_with_guide(
        "비교 분석", "compare_input", "공통점과 차이점", [
            "기준(관점, 목적 등)을 설정해 비교",
            "논리적으로 유사점·차이점 제시"
        ],
        hint_key="compare", hint_prompt="두 기사 간 공통점과 차이점을 비교하여 분석하는 문단을 위한 문장 구성 힌트를 제시해줘. (한국어)"
    )

    conclusion = paragraph_input_with_guide(
        "결론", "conclusion_input", "요약 및 의견", [
            "전체 비교 내용 요약",
            "자신의 의견이나 평가 포함"
        ],
        hint_key="conclusion", hint_prompt="비교 설명문 결론에 사용할 수 있는 마무리 문장 3개를 제안해줘. (한국어)"
    )

    st.markdown("---")
    st.markdown("### 전체 초안 미리보기")

    full_draft = "\n\n".join([
        f"[서론]\n{st.session_state.intro_input}",
        f"[본론 1 - 기사 1]\n{st.session_state.body1_input}",
        f"[본론 2 - 기사 2]\n{st.session_state.body2_input}",
        f"[비교 분석]\n{st.session_state.compare_input}",
        f"[결론]\n{st.session_state.conclusion_input}"
    ])

    st.session_state.draft = full_draft

    st.markdown(f"""<div style="background-color:#f9f9f9; padding:15px; border-radius:10px; color:black; font-size:16px;">
<pre style="white-space: pre-wrap; word-wrap: break-word;">{full_draft}</pre>
</div>""", unsafe_allow_html=True)

    if st.session_state.paragraph_feedback:
        st.markdown("---")
        st.markdown("### 문단별 피드백 요약")
        summary_text = summarize_paragraph_feedback(st.session_state.paragraph_feedback)
        st.info(summary_text)

    col_btn1, col_btn2 = st.columns([1, 1])
    with col_btn1:
        if st.button("← 이전 단계", use_container_width=True):
            st.session_state.stage = "analysis"
            st.rerun()

    with col_btn2:
        overall_draft_error = st.empty()
        if st.button("AI 피드백 받기 →", type="primary", use_container_width=True):
            paragraphs = [st.session_state.intro_input, st.session_state.body1_input,
                          st.session_state.body2_input, st.session_state.compare_input,
                          st.session_state.conclusion_input]
            
            is_valid = all(p.strip() for p in paragraphs)
            
            if is_valid:
                overall_draft_error.empty()
                st.session_state.stage = "feedback"
                st.rerun()
            else:
                overall_draft_error.error("모든 문단을 작성해주세요.")

elif st.session_state.stage == "feedback":
    st.subheader("4단계. AI 피드백")
    
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.markdown("**내 초안**")
        st.text_area(
            "작성한 초안",
            value=st.session_state.draft,
            height=400,
            disabled=True,
            key="draft_display_feedback"
        )
    
    with col2:
        st.markdown("**AI 피드백**")
        
        # AI 피드백 생성 (한 번만)
        if not st.session_state.get("feedback"):
            if OPENAI_OK:
                with st.spinner("AI 피드백을 생성하고 있습니다..."):
                    st.info("종합 피드백 생성 중...", icon="🧠")
                    feedback = gpt_feedback(st.session_state.draft)
                    st.session_state.feedback = feedback
                    
                    st.info("영어 표현 능력 평가 중...", icon="📝")
                    english_draft = translate_to_english(st.session_state.draft)
                    st.session_state.writing_evaluation = evaluate_writing_rubric(english_draft)
            else:
                st.session_state.feedback = "AI 피드백 기능이 비활성화되어 있습니다."
        
        tab1, tab2 = st.tabs(["AI 피드백", "루브릭 평가"])
        
        with tab1:
            st.text_area(
                "AI 피드백",
                value=st.session_state.feedback,
                height=400,
                disabled=True
            )
        
        with tab2:
            st.markdown("#### 영어 표현 능력 평가")
            eval_data = format_analysis_for_display(st.session_state.writing_evaluation, "evaluation")
            
            if "error" not in eval_data:
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    if "내용논리성" in eval_data and isinstance(eval_data["내용논리성"], dict):
                        logic = eval_data["내용논리성"]
                        st.metric("내용 논리성", f"{logic.get('점수', 0)}/4점")
                        st.caption(logic.get('근거', ''))
                    else:
                        st.metric("내용 논리성", "N/A")
                
                with col2:
                    if "구성체계성" in eval_data and isinstance(eval_data["구성체계성"], dict):
                        org = eval_data["구성체계성"]
                        st.metric("구성 체계성", f"{org.get('점수', 0)}/4점")
                        st.caption(org.get('근거', ''))
                    else:
                        st.metric("구성 체계성", "N/A")
                
                with col3:
                    if "문법어휘정확성" in eval_data and isinstance(eval_data["문법어휘정확성"], dict):
                        lang = eval_data["문법어휘정확성"]
                        st.metric("문법·어휘", f"{lang.get('점수', 0)}/4점")
                        st.caption(lang.get('근거', ''))
                    else:
                        st.metric("문법·어휘", "N/A")
                
                if eval_data.get('총점'):
                    st.markdown(f"**총점**: {eval_data['총점']}")
                
                if eval_data.get('종합평가'):
                    st.markdown("**종합 평가**")
                    st.info(eval_data['종합평가'])
            else:
                st.error(eval_data.get("error", "평가 오류"))
    
    st.markdown("---")
    st.markdown("#### 피드백 성찰")
    feedback_reflection_error = st.empty()
    feedback_reflection = st.text_area(
        "AI 피드백을 받은 후 느낀 점과 개선하고 싶은 부분을 적어보세요:",
        height=100,
        key="feedback_reflection"
    )
    
    col_btn1, col_btn2 = st.columns([1, 1])
    with col_btn1:
        if st.button("← 이전 단계", use_container_width=True):
            st.session_state.stage = "draft"
            st.rerun()
    
    with col_btn2:
        if st.button("최종 수정 단계로 →", type="primary", use_container_width=True):
            if not feedback_reflection.strip():
                feedback_reflection_error.error("피드백 성찰 내용을 입력해주세요.")
            else:
                feedback_reflection_error.empty()
                st.session_state.reflection_log.append({
                    "stage": "feedback",
                    "content": feedback_reflection,
                    "timestamp": datetime.datetime.now()
                })
                with st.spinner("성찰 내용을 바탕으로 문제해결 역량 평가 중..."):
                    st.session_state.problem_solving_score = assess_problem_solving(feedback_reflection)
                st.session_state.stage = "final"
                st.rerun()

elif st.session_state.stage == "final":
    st.subheader("5단계. 최종 수정 및 완성")
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.markdown("**최종 수정**")
        if "final_text" not in st.session_state or not st.session_state.final_text:
             st.session_state.final_text = st.session_state.draft
             
        final_text = st.text_area(
            "피드백을 반영하여 최종 수정하세요",
            value=st.session_state.final_text,
            height=400,
            key="final_input"
        )
        
        st.session_state.final_text = final_text
    
    with col2:
        st.markdown("**종합 평가 결과**")
        
        if st.session_state.problem_solving_score:
            with st.expander("문제해결 역량 평가", expanded=True):
                problem_data = format_analysis_for_display(st.session_state.problem_solving_score, "assessment")
                
                if "error" not in problem_data:
                    if "assessment" in problem_data:
                        st.info(problem_data["assessment"])
                    else:
                        col1, col2 = st.columns(2)
                        
                        areas = ["문제이해", "분석적사고", "대안발견및기획", "의사소통"]
                        for i, area in enumerate(areas):
                            col = col1 if i % 2 == 0 else col2
                            with col:
                                if area in problem_data and isinstance(problem_data[area], dict):
                                    score = problem_data[area].get('점수', 0)
                                    st.metric(area.replace('및', ' & '), f"{score}/5점")
                                elif area in problem_data:
                                    st.metric(area.replace('및', ' & '), str(problem_data[area]))
                        
                        if problem_data.get('총점'):
                            st.markdown(f"**총점**: {problem_data['총점']}")
                        
                        if problem_data.get('종합평가'):
                            st.markdown("**종합 평가**")
                            st.info(problem_data['종합평가'])
                else:
                    st.error(problem_data.get("error", "평가 오류"))
        
        if st.session_state.writing_evaluation:
            with st.expander("영어 표현 능력 평가", expanded=True):
                eval_data = format_analysis_for_display(st.session_state.writing_evaluation, "evaluation")
                
                if "error" not in eval_data:
                    col1, col2, col3 = st.columns(3)
                    
                    with col1:
                        if "내용논리성" in eval_data and isinstance(eval_data["내용논리성"], dict):
                            logic = eval_data["내용논리성"]
                            score = logic.get('점수', 0)
                            st.metric("내용 논리성", f"{score}/4점")
                        else:
                            st.metric("내용 논리성", "N/A")
                    
                    with col2:
                        if "구성체계성" in eval_data and isinstance(eval_data["구성체계성"], dict):
                            org = eval_data["구성체계성"]
                            score = org.get('점수', 0)
                            st.metric("구성 체계성", f"{score}/4점")
                        else:
                            st.metric("구성 체계성", "N/A")
                    
                    with col3:
                        if "문법어휘정확성" in eval_data and isinstance(eval_data["문법어휘정확성"], dict):
                            lang = eval_data["문법어휘정확성"]
                            score = lang.get('점수', 0)
                            st.metric("문법·어휘", f"{score}/4점")
                        else:
                            st.metric("문법·어휘", "N/A")
                else:
                    st.error(eval_data.get("error", "평가 오류"))
    
    st.markdown("---")
    
    col_btn1, col_btn2, col_btn3, col_btn4 = st.columns([1, 1, 1, 1])
    
    with col_btn1:
        if st.button("← 이전 단계", use_container_width=True):
            st.session_state.stage = "feedback"
            st.rerun()

    with col_btn2:
        analysis_summary = {
            "논조분석1": st.session_state.tone_analysis1,
            "논조분석2": st.session_state.tone_analysis2,
            "영어표현평가": st.session_state.writing_evaluation,
            "문단별피드백": summarize_paragraph_feedback(st.session_state.paragraph_feedback)
        }
        if st.session_state.problem_solving_score:
            analysis_summary["문제해결평가"] = st.session_state.problem_solving_score
            
        docx_data = create_docx_content(final_text, analysis_summary)
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"news_comparison_complete_{timestamp}.docx"
        st.download_button(
            label="종합 보고서 다운로드",
            data=docx_data,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            disabled=not final_text.strip()
        )

    with col_btn4:
        if st.button("처음부터 다시", use_container_width=True):
            for key in list(st.session_state.keys()):
                if key != 'stage':
                    st.session_state.pop(key)
            st.session_state.stage = "input"
            st.rerun()
    
    if final_text.strip():
        st.markdown("### 완성된 작문 미리보기")
        st.success(final_text)
        
        if st.session_state.reflection_log:
            with st.expander("학습 성찰 기록", expanded=False):
                for idx, log in enumerate(st.session_state.reflection_log):
                    st.markdown(f"**{log['stage']} 단계 성찰:**")
                    st.write(log['content'])
                    st.caption(f"작성 시간: {log['timestamp']}")
                    st.markdown("---")

with st.sidebar:
    st.markdown("### 사용 방법")
    st.markdown("""
    1. **기사 입력**: 비교할 두 기사의 본문을 파일 업로드
    2. **논조 분석**: AI가 각 기사의 논조와 입장을 분석
    3. **초안 작성**: 분석 결과를 참고하여 비교 설명문 작성
       - 각 문단 완료 시 즉시 피드백 가능
    4. **AI 피드백**: 루브릭 기반 AI 피드백
    5. **최종 완성**: 피드백을 반영한 수정 후 종합 보고서 다운로드
    """)
    
    st.markdown("### 설정 상태")
    if OPENAI_OK:
        st.success("OpenAI API 연결됨")
    else:
        st.error("OpenAI API 연결 실패")
    
    st.markdown("### 진행 상황")
    st.markdown(f"현재 단계: **{stage_names[current_stage_idx]}**")
    
    checklist_items = [
        ("기사 입력", bool(st.session_state.get("article1") and st.session_state.get("article2"))),
        ("논조 분석", bool(st.session_state.get("tone_analysis1") and st.session_state.get("tone_analysis2"))),
        ("초안 작성", bool(st.session_state.get("draft"))),
        ("문단별 피드백", bool(st.session_state.get("paragraph_feedback"))),
        ("AI 피드백", bool(st.session_state.get("feedback"))),
        ("루브릭 평가", bool(st.session_state.get("writing_evaluation"))),
        ("최종 완성", bool(st.session_state.get("final_text")))
    ]
    
    for item, completed in checklist_items:
        icon = "✅" if completed else "⏳"
        st.markdown(f"{icon} {item}")